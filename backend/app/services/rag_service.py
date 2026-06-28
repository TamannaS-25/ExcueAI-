import os
import re
import json
import numpy as np
from typing import List, Dict, Any, Optional
from backend.app.database import get_db

# Try to import pypdf for PDF processing
try:
    import pypdf
except ImportError:
    pypdf = None

# Try to import python-docx
try:
    import docx
except ImportError:
    docx = None

# Try to import openai
try:
    import openai
    from openai import OpenAI
except ImportError:
    openai = None
    OpenAI = None

# Simple custom vector store to act as ChromaDB fallback
class SimpleVectorStore:
    def __init__(self, filepath: str = "excueai_vector_store.json"):
        self.filepath = filepath
        self.documents = []
        self.load()

    def load(self):
        if os.path.exists(self.filepath):
            try:
                with open(self.filepath, "r", encoding="utf-8") as f:
                    self.documents = json.load(f)
            except Exception as e:
                print(f"Error loading vector store: {e}")
                self.documents = []

    def save(self):
        try:
            with open(self.filepath, "w", encoding="utf-8") as f:
                json.dump(self.documents, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving vector store: {e}")

    def get_simple_embedding(self, text: str) -> List[float]:
        # Generate a deterministic mock TF-IDF-like embedding based on word frequencies
        # Vocabulary size = 128 dimensions
        words = re.findall(r'\w+', text.lower())
        vector = [0.0] * 128
        for word in words:
            # Hash word to a position in the 128-dim vector
            idx = sum(ord(c) for c in word) % 128
            vector[idx] += 1.0
        
        # L2 Normalize the vector
        norm = np.linalg.norm(vector)
        if norm > 0:
            vector = [float(v / norm) for v in vector]
        return vector

    def get_embedding(self, text: str, api_key: Optional[str] = None) -> List[float]:
        if api_key and openai:
            try:
                client = OpenAI(api_key=api_key)
                response = client.embeddings.create(
                    input=[text],
                    model="text-embedding-3-small"
                )
                return response.data[0].embedding
            except Exception as e:
                print(f"OpenAI embedding generation failed, falling back to simple vector: {e}")
        
        return self.get_simple_embedding(text)

    def add_documents(self, texts: List[str], metadatas: List[Dict[str, Any]], ids: List[str], api_key: Optional[str] = None):
        for text, meta, doc_id in zip(texts, metadatas, ids):
            embedding = self.get_embedding(text, api_key)
            self.documents.append({
                "id": doc_id,
                "text": text,
                "metadata": meta,
                "embedding": embedding
            })
        self.save()

    def query(self, query_text: str, k: int = 3, api_key: Optional[str] = None) -> List[Dict[str, Any]]:
        query_vector = np.array(self.get_embedding(query_text, api_key))
        
        results = []
        for doc in self.documents:
            doc_vector = np.array(doc["embedding"])
            # Compute cosine similarity
            if len(query_vector) == len(doc_vector):
                similarity = float(np.dot(query_vector, doc_vector) / (np.linalg.norm(query_vector) * np.linalg.norm(doc_vector) + 1e-9))
                results.append((similarity, doc))
        
        # Sort by similarity descending
        results.sort(key=lambda x: x[0], reverse=True)
        
        return [
            {
                "text": item[1]["text"],
                "metadata": item[1]["metadata"],
                "similarity": item[0]
            }
            for item in results[:k]
        ]

    def delete_by_document_id(self, document_id: int):
        self.documents = [doc for doc in self.documents if doc["metadata"].get("document_id") != document_id]
        self.save()

# Instantiate global fallback vector store
vector_store = SimpleVectorStore()

def extract_text_from_pdf(filepath: str) -> str:
    text = ""
    if not pypdf:
        # Fallback if pypdf is not installed: try to read plain text
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except:
            return "[PDF reader not installed. Could not extract text content.]"
            
    try:
        reader = pypdf.PdfReader(filepath)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
    return text

def extract_text_from_docx(filepath: str) -> str:
    text = ""
    if not docx:
        return "[DOCX reader not installed. Could not extract text content.]"
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
    return text

def extract_text_from_file(filepath: str) -> str:
    _, ext = os.path.splitext(filepath.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:  # TXT or similar
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file {filepath}: {e}")
            return ""

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    # Split text into overlapping chunks
    text = re.sub(r'\s+', ' ', text).strip()
    if len(text) <= chunk_size:
        return [text]
        
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start += chunk_size - chunk_overlap
    return chunks

def index_document(document_id: int, filename: str, filepath: str, api_key: Optional[str] = None):
    # Extract
    text = extract_text_from_file(filepath)
    if not text:
        print(f"No text extracted from document {filename}")
        return
        
    # Chunk
    chunks = chunk_text(text)
    
    # Store
    texts = []
    metadatas = []
    ids = []
    
    for i, chunk in enumerate(chunks):
        texts.append(chunk)
        metadatas.append({
            "document_id": document_id,
            "filename": filename,
            "chunk_index": i
        })
        ids.append(f"doc_{document_id}_chunk_{i}")
        
    vector_store.add_documents(texts, metadatas, ids, api_key=api_key)

def delete_document_index(document_id: int):
    vector_store.delete_by_document_id(document_id)

def query_knowledge_base(query_text: str, k: int = 3, api_key: Optional[str] = None) -> List[Dict[str, Any]]:
    return vector_store.query(query_text, k=k, api_key=api_key)
